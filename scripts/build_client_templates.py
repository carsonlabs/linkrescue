from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "client-templates"
OUT.mkdir(parents=True, exist_ok=True)

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
GREEN = "1F6B4E"
GOLD = "7A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"


def set_run(run, size=11, color=NAVY, bold=False, italic=False, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_borders(table, color="D9E2EC", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_para(paragraph, before=0, after=6, line=1.1, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep


def add_text(doc, text, style="Normal", before=0, after=6, bold_lead=None, align=None):
    p = doc.add_paragraph(style=style)
    set_para(p, before, after, 1.1 if style == "Normal" else 1.0, style.startswith("Heading"))
    if align:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run(lead, 11, NAVY, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run(rest, 11, NAVY)
    else:
        run = p.add_run(text)
        size = 11
        color = NAVY
        if style == "Heading 1":
            size, color = 16, BLUE
        elif style == "Heading 2":
            size, color = 13, BLUE
        elif style == "Heading 3":
            size, color = 12, DARK_BLUE
        set_run(run, size, color, bold=style.startswith("Heading"))
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_para(p, 0, 4, 1.208)
    run = p.add_run(text)
    set_run(run)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_para(p, 0, 4, 1.208)
    run = p.add_run(text)
    set_run(run)
    return p


def add_callout(doc, label, text, color=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_borders(table, "C8D8EA", "8")
    cell = table.cell(0, 0)
    shade(cell, color)
    p = cell.paragraphs[0]
    set_para(p, 0, 0, 1.1)
    label_run = p.add_run(f"{label}  ")
    set_run(label_run, 10, DARK_BLUE, bold=True)
    body = p.add_run(text)
    set_run(body, 10.5, NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_metadata(doc, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    set_table_geometry(table, [2700, 6660])
    set_borders(table, "D9E2EC", "4")
    set_repeat_table_header(table.rows[0])
    for cell, value in zip(table.rows[0].cells, ("Field", "Details")):
        shade(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        set_para(p, 0, 0)
        r = p.add_run(value)
        set_run(r, 9.5, DARK_BLUE, bold=True)
    for idx, (label, value) in enumerate(rows):
        left, right = table.rows[idx + 1].cells
        shade(left, LIGHT_GRAY)
        p = left.paragraphs[0]
        set_para(p, 0, 0)
        r = p.add_run(label)
        set_run(r, 10, MUTED, bold=True)
        p = right.paragraphs[0]
        set_para(p, 0, 0)
        r = p.add_run(value)
        set_run(r, 10.5, NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_two_column_metadata(doc, left_rows, right_rows):
    table = doc.add_table(rows=max(len(left_rows), len(right_rows)) + 1, cols=4)
    set_table_geometry(table, [1350, 3330, 1350, 3330])
    set_borders(table, "D9E2EC", "4")
    set_repeat_table_header(table.rows[0])
    for cell, value in zip(table.rows[0].cells, ("Field", "Details", "Field", "Details")):
        shade(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        set_para(p, 0, 0)
        r = p.add_run(value)
        set_run(r, 9.5, DARK_BLUE, bold=True)
    for row_index in range(1, len(table.rows)):
        source_index = row_index - 1
        values = list(left_rows[source_index]) if source_index < len(left_rows) else ["", ""]
        values += list(right_rows[source_index]) if source_index < len(right_rows) else ["", ""]
        for cell_index, value in enumerate(values):
            cell = table.rows[row_index].cells[cell_index]
            if cell_index in (0, 2):
                shade(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            set_para(p, 0, 0)
            r = p.add_run(value)
            set_run(r, 9.5 if cell_index in (0, 2) else 10, MUTED if cell_index in (0, 2) else NAVY, bold=cell_index in (0, 2))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_borders(table)
    set_repeat_table_header(table.rows[0])
    for cell, header in zip(table.rows[0].cells, headers):
        shade(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        set_para(p, 0, 0)
        r = p.add_run(header)
        set_run(r, 9.5, DARK_BLUE, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            p = cell.paragraphs[0]
            set_para(p, 0, 0)
            r = p.add_run(value)
            set_run(r, 10, NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_checkbox_line(doc, label, placeholder=""):
    p = doc.add_paragraph()
    set_para(p, 0, 5, 1.1)
    r = p.add_run("[ ]  ")
    set_run(r, 11, DARK_BLUE, bold=True)
    r = p.add_run(label)
    set_run(r, 10.5, NAVY, bold=True)
    if placeholder:
        r = p.add_run(f"  {placeholder}")
        set_run(r, 10, MUTED, italic=True)
    return p


def base_document(header_label):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para(header, 0, 0, 1.0)
    r = header.add_run("LINKRESCUE  |  ")
    set_run(r, 8.5, BLUE, bold=True)
    r = header.add_run(header_label)
    set_run(r, 8.5, MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para(footer, 0, 0, 1.0)
    r = footer.add_run("LinkRescue | Confidential client working document")
    set_run(r, 8.5, MUTED)
    return doc


def add_customer_title(doc, kicker, title, subtitle):
    p = doc.add_paragraph()
    set_para(p, 0, 2, 1.0)
    r = p.add_run(kicker.upper())
    set_run(r, 10, BLUE, bold=True)
    p = doc.add_paragraph()
    set_para(p, 0, 8, 1.0)
    r = p.add_run(title)
    set_run(r, 29, NAVY, bold=True)
    p = doc.add_paragraph()
    set_para(p, 0, 18, 1.1)
    r = p.add_run(subtitle)
    set_run(r, 13, MUTED)


def build_proposal():
    doc = base_document("Recovery Sprint Proposal")
    add_customer_title(
        doc,
        "Fixed-scope service proposal",
        "Recovery Sprint",
        "A focused technical review for an affiliate content site.",
    )
    add_two_column_metadata(
        doc,
        [("Prepared for", "[Client name]"), ("Website", "[client-site.com]"), ("Date", str(date.today().strftime("%B %d, %Y")))],
        [("Prepared by", "LinkRescue"), ("Engagement", "Fixed-scope Recovery Sprint"), ("Investment", "US $499")],
    )
    add_callout(
        doc,
        "Purpose",
        "This proposal turns observable link evidence into a prioritized technical repair plan. It does not estimate lost revenue, guarantee commissions, or authorize work beyond the agreed scope.",
    )

    add_text(doc, "1. The engagement", "Heading 1")
    add_text(doc, "LinkRescue will review the agreed public pages and outbound links for technical issues that can be observed without access to analytics, affiliate dashboards, or private systems. The result is a clear, written repair list for the client's team or developer.")

    add_text(doc, "2. Scope of work", "Heading 1")
    add_table(doc, ["Included", "What LinkRescue will deliver"], [
        ("Public-page review", "Review up to [20] agreed public pages or the agreed crawl sample."),
        ("Outbound-link checks", "Identify observable broken links, non-working destinations, redirect behavior, and visible tracking-parameter risks."),
        ("Prioritization", "Rank findings by technical severity, affected pages, and practical repair sequence."),
        ("Implementation notes", "Provide concise notes for the client or developer to repair, replace, redirect, or verify each priority item."),
        ("Readout", "One written report plus one asynchronous clarification round within five business days of delivery."),
    ], [2600, 6760])

    add_text(doc, "3. Not included", "Heading 1")
    for text in [
        "Revenue, commission, traffic, ranking, or conversion forecasts or guarantees.",
        "Changes to the client's website, affiliate accounts, merchant accounts, analytics, or hosting.",
        "Private, logged-in, paywalled, blocked, or inaccessible pages and destinations.",
        "Legal, tax, merchant-program, privacy, security, or compliance advice.",
        "Ongoing monitoring, emergency support, or work beyond the written scope.",
    ]:
        add_bullet(doc, text)

    add_text(doc, "4. Client inputs needed", "Heading 1")
    for text in [
        "The primary site URL and up to [20] priority pages, if applicable.",
        "The affiliate programmes or merchants that matter most.",
        "Known problem links, recent merchant changes, or priority commercial pages.",
        "The primary audience country or market, where this affects merchant destinations.",
    ]:
        add_bullet(doc, text)

    add_text(doc, "5. Timing and delivery", "Heading 1")
    add_table(doc, ["Milestone", "Timing"], [
        ("Scope confirmed", "After written acceptance and the required client inputs."),
        ("Review begins", "Within two business days after the scope is confirmed."),
        ("Written delivery", "Within five business days after the review begins."),
        ("Clarification round", "One asynchronous round within five business days of the written delivery."),
    ], [3200, 6160])

    add_text(doc, "6. Investment and acceptance", "Heading 1")
    add_text(doc, "Fixed fee: US $499. Payment instructions will be issued only after the client confirms this scope in writing. Work begins after scope and payment timing are agreed. Either party may pause before work begins if the site is not a fit or the scope cannot be confirmed.")
    add_callout(doc, "To accept", "Reply by email confirming: “I accept the Recovery Sprint scope for [site] at US $499.” LinkRescue will then send the practical next steps.", "EAF4EF")

    add_text(doc, "7. Plain-language terms", "Heading 1")
    for text in [
        "The client remains responsible for all implementation decisions and changes to its site or affiliate accounts.",
        "LinkRescue will treat client-provided context as confidential and will not publish it without written permission.",
        "The review is a technical assessment of observable conditions at the time checked. External sites and programmes can change after delivery.",
        "Any material change to scope, pages, timing, or fee requires written agreement before work continues.",
    ]:
        add_bullet(doc, text)

    add_text(doc, "Acceptance record", "Heading 2")
    add_checkbox_line(doc, "Client confirms the agreed website and priority pages", "[name / date]")
    add_checkbox_line(doc, "Client accepts the fixed-scope Recovery Sprint at US $499", "[name / date]")
    add_checkbox_line(doc, "Client has read the exclusions and plain-language terms above", "[name / date]")
    add_text(doc, "This proposal is a practical service-scope template, not legal advice. Obtain local legal advice before adapting it into a contract.", before=8, after=0)
    doc.save(OUT / "LinkRescue_Recovery_Sprint_Proposal_Template.docx")


def build_delivery_report():
    doc = base_document("Recovery Sprint Delivery Report")
    add_customer_title(
        doc,
        "Client delivery report",
        "Recovery Sprint Report",
        "Prioritized technical link findings and implementation notes.",
    )
    add_two_column_metadata(
        doc,
        [("Client", "[Client name]"), ("Website", "[client-site.com]"), ("Review period", "[date range]")],
        [("Prepared by", "LinkRescue"), ("Report date", "[date]"), ("Scope reference", "[proposal / email date]")],
    )
    add_callout(
        doc,
        "Read this first",
        "This report records observable technical conditions at the time of review. It is not a revenue forecast or a guarantee of commissions, traffic, ranking, merchant acceptance, or future link behavior.",
        "FFF7E8",
    )

    add_text(doc, "1. Executive summary", "Heading 1")
    add_metadata(doc, [
        ("Pages reviewed", "[number / list or sample description]"),
        ("Links checked", "[number]"),
        ("Priority findings", "[number]"),
        ("Immediate action", "[one sentence on the highest-priority repair]"),
    ])
    add_text(doc, "Summary for the client: [Write 3-5 sentences on what was observed, which category matters most, and the recommended order of action. Keep claims technical and specific.]", before=6, after=4)

    add_text(doc, "2. Priority repair queue", "Heading 1")
    add_text(doc, "Replace each sample row with an observed finding. Use one row per actionable item; attach or link to evidence only where you have permission to share it.", before=0, after=6)
    add_table(doc, ["Priority", "Page / link", "Observed condition", "Recommended action", "Owner"], [
        ("P1", "[page or link]", "[e.g., destination returns a visible 404]", "[replace, update, or remove link]", "[client / developer]"),
        ("P1", "[page or link]", "[e.g., redirect lands on generic homepage]", "[confirm destination; update link if necessary]", "[client / developer]"),
        ("P2", "[page or link]", "[e.g., visible tracking parameter missing after redirect]", "[verify affiliate URL and update if appropriate]", "[client / affiliate manager]"),
        ("P3", "[page or link]", "[e.g., slow or inconsistent external response]", "[retest; replace if issue persists]", "[client]"),
    ], [700, 1900, 2500, 2860, 1400])

    add_text(doc, "3. Evidence and implementation notes", "Heading 1")
    add_text(doc, "For each P1 or P2 item, add concise detail using the pattern below. Do not infer why a merchant changed a page or whether a specific commission was lost unless the client supplies independent evidence.")
    add_text(doc, "Finding [#]: [Short title]", "Heading 2")
    add_metadata(doc, [
        ("Affected page", "[URL]"),
        ("Outbound destination", "[URL or merchant]"),
        ("Observed", "[status, redirect path, parameter observation, or technical evidence]"),
        ("Recommended repair", "[specific next step]"),
        ("Verification", "[how the client can confirm the repair]"),
    ])
    add_text(doc, "Notes: [Add anything the implementer needs, keeping evidence factual and time-bound.]", before=6, after=4)

    add_text(doc, "4. Recommended order of work", "Heading 1")
    for text in [
        "Repair or remove P1 destinations that are plainly unavailable or misdirected.",
        "Verify P1/P2 affiliate URLs with the relevant merchant or programme where needed.",
        "Update the affected page copy and links, then test the public destination again.",
        "Keep a simple change log: page, old URL, new URL, date, and person responsible.",
        "Re-check a small sample after implementation; external destinations may change over time.",
    ]:
        add_number(doc, text)

    add_text(doc, "5. Completion checklist", "Heading 1")
    add_checkbox_line(doc, "Client has assigned owners for all P1 findings")
    add_checkbox_line(doc, "Client has completed or scheduled priority link repairs")
    add_checkbox_line(doc, "Client has re-tested public destinations after changes")
    add_checkbox_line(doc, "Client has stored the repair log with this report")

    add_text(doc, "6. Optional next step", "Heading 1")
    add_text(doc, "If ongoing checks would be useful after the repair work is complete, LinkRescue can discuss whether the Monitoring Desk is a fit. Availability, scope, and price are confirmed separately; no monitoring is started automatically.")
    add_text(doc, "This template is a practical reporting aid, not legal advice. Keep client-specific information confidential and obtain appropriate advice for contractual, privacy, or regulatory requirements.", before=10, after=0)
    doc.save(OUT / "LinkRescue_Recovery_Sprint_Delivery_Report_Template.docx")


if __name__ == "__main__":
    build_proposal()
    build_delivery_report()
